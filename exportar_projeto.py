import os
import zipfile
import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from reportlab.lib.pagesizes import landscape, A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from datetime import datetime
import threading
import time

PASTA_BASE = r"C:\Users\User\PycharmProjects"
EXTENSOES_VALIDAS = ('.py', '.html', '.css', '.js', '.json', '.txt', '.md', '.jpg', '.jpeg', '.png')

secoes = {
    'Arquivos Python': ['.py'],
    'Arquivos HTML': ['.html'],
    'Arquivos CSS': ['.css'],
    'Arquivos JavaScript': ['.js'],
    'Outros Arquivos de Texto': ['.json', '.txt', '.md']
}

def exportar_projeto(nome_projeto, progressbar):
    try:
        pasta_projeto = os.path.join(PASTA_BASE, nome_projeto)
        pasta_saida = os.path.join(pasta_projeto, 'exportacao')
        os.makedirs(pasta_saida, exist_ok=True)
        timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')

        nome_zip = os.path.join(pasta_saida, f'{nome_projeto}_exportado_{timestamp}.zip')
        nome_docx = os.path.join(pasta_saida, f'{nome_projeto}_documentado_{timestamp}.docx')
        nome_pdf = os.path.join(pasta_saida, f'{nome_projeto}_documentado_{timestamp}.pdf')

        arquivos_para_processar = []

        # Coletar todos os arquivos válidos
        for pasta_atual, _, arquivos in os.walk(pasta_projeto):
            if 'venv' in pasta_atual or '__pycache__' in pasta_atual or 'exportacao' in pasta_atual:
                continue
            for arquivo in arquivos:
                if arquivo.endswith(EXTENSOES_VALIDAS):
                    arquivos_para_processar.append((pasta_atual, arquivo))

        total = len(arquivos_para_processar)
        progresso = 0

        # Criar ZIP
        with zipfile.ZipFile(nome_zip, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for pasta_atual, arquivo in arquivos_para_processar:
                caminho = os.path.join(pasta_atual, arquivo)
                rel = os.path.relpath(caminho, pasta_projeto)
                zipf.write(caminho, rel)

        # Criar DOCX e PDF
        documento = Document()
        documento.add_heading(f'Documentação do Projeto: {nome_projeto}', level=1)
        documento.add_paragraph(f'Exportado em: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}\n')

        pdf = canvas.Canvas(nome_pdf, pagesize=landscape(A4))
        largura, altura = landscape(A4)
        x_ini, y_ini = 2 * cm, altura - 2 * cm
        linha_altura = 12
        y = y_ini

        pdf.setFont("Helvetica-Bold", 14)
        pdf.drawString(x_ini, y, f'Documentação do Projeto: {nome_projeto}')
        y -= 2 * linha_altura
        pdf.setFont("Courier", 10)

        for nome_secao, extensoes in secoes.items():
            documento.add_heading(nome_secao, level=1)
            pdf.setFont("Helvetica-Bold", 12)
            if y <= 3 * cm:
                pdf.showPage()
                y = y_ini
            pdf.drawString(x_ini, y, f'### {nome_secao} ###')
            y -= 2 * linha_altura

            for pasta_atual, arquivo in arquivos_para_processar:
                if any(arquivo.endswith(ext) for ext in extensoes):
                    caminho = os.path.join(pasta_atual, arquivo)
                    rel = os.path.relpath(caminho, pasta_projeto)

                    documento.add_heading(rel, level=2)
                    try:
                        with open(caminho, 'r', encoding='utf-8') as f:
                            conteudo = f.read()
                            documento.add_paragraph(conteudo)

                            if y <= 3 * cm:
                                pdf.showPage()
                                y = y_ini
                                pdf.setFont("Helvetica-Bold", 12)
                                pdf.drawString(x_ini, y, f'### {nome_secao} ###')
                                y -= 2 * linha_altura
                                pdf.setFont("Courier", 9)

                            pdf.setFont("Courier-Bold", 9)
                            pdf.drawString(x_ini, y, f'Arquivo: {rel}')
                            y -= linha_altura
                            pdf.setFont("Courier", 8)

                            for linha in conteudo.split('\n'):
                                if y <= 2 * cm:
                                    pdf.showPage()
                                    y = y_ini
                                pdf.drawString(x_ini, y, linha[:200])
                                y -= linha_altura
                            y -= linha_altura
                    except Exception as e:
                        documento.add_paragraph(f'[Erro: {e}]')

                progresso += 1
                progress_percent = int((progresso / total) * 100)
                progressbar['value'] = progress_percent
                root.update_idletasks()

        documento.save(nome_docx)
        pdf.save()

        progressbar['value'] = 100
        messagebox.showinfo("Concluído", f"Exportação finalizada com sucesso em:\n{pasta_saida}")
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro: {e}")

# INTERFACE
root = tk.Tk()
root.title("Exportador de Projeto")
root.geometry("500x250")

ttk.Label(root, text="Selecione o projeto para exportar:").pack(pady=10)

projetos = [p for p in os.listdir(PASTA_BASE) if os.path.isdir(os.path.join(PASTA_BASE, p))]
projeto_var = tk.StringVar()
combo = ttk.Combobox(root, textvariable=projeto_var, values=projetos, state='readonly', width=60)
combo.pack(pady=5)

progressbar = ttk.Progressbar(root, orient="horizontal", length=400, mode="determinate")
progressbar.pack(pady=20)

def ao_exportar():
    projeto = projeto_var.get()
    if projeto:
        progressbar['value'] = 0
        thread = threading.Thread(target=exportar_projeto, args=(projeto, progressbar))
        thread.start()
    else:
        messagebox.showwarning("Seleção necessária", "Escolha um projeto antes de exportar.")

ttk.Button(root, text="Exportar", command=ao_exportar).pack(pady=10)

root.mainloop()
